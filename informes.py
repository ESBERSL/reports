import pandas as pd
from docx import Document
from io import BytesIO
from supabase import create_client, Client
import streamlit as st
from datetime import datetime
from database import  obtener_defectos, obtener_datos_cuadro, obtener_cuadros
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from collections import defaultdict
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import uuid 
from docx.shared import Inches

# Conexión a Supabase
url = st.secrets["supabase"]["SUPABASE_URL"]
key = st.secrets["supabase"]["SUPABASE_KEY"]
supabase: Client = create_client(url, key)


def obtener_datos_centro(centro_id, tabla='centros'):
    response = supabase.table(tabla).select('*').eq('id', centro_id).execute()
    
    if response.data:
        return response.data[0]  # Retorna el primer resultado
    else:
        return {"nombre": "Desconocido", "direccion": "Desconocida"}

# Ruta de la plantilla de Word (asegúrate de que exista en tu proyecto)
PLANTILLA_TIERRAS = "tierras_plantilla.docx"
PLANTILLA_AISLAMIENTOS = "aislamientos_plantilla.docx"

    
# Función para modificar la plantilla de Word
def generar_informe_word_tierras(centro_id):
    # Cargar la plantilla de Word
    doc = Document(PLANTILLA_TIERRAS)
    datos_centro = obtener_datos_centro(centro_id)
    nombre_centro = datos_centro.get("nombre") or "Desconocido"
    direccion_centro = datos_centro.get("direccion") or "Desconocido"
    cp_centro = str(int(float(datos_centro.get("cp")))) if datos_centro.get("cp") else "00000"
    provincia_centro = datos_centro.get("provincia") or "Desconocido"
    pueblo_centro = datos_centro.get("pueblo") or "Desconocido"
    fecha_actual = datetime.now()



    df_cuadros = obtener_cuadros(centro_id)
    medidas_tierra = df_cuadros["tierra_ohmnios"].tolist()

    # Analizar la medición más alta y comparar con 48
    medida_maxima = max(medidas_tierra)
    informe_estado = "Favorable" if medida_maxima <= 48 else "Desfavorable"

    reemplazos = {
        "[NOMBRE]": nombre_centro,
        "[DIRECCION]": direccion_centro,
        "[CP]": cp_centro,
        "[PROVINCIA]": provincia_centro,
        "[PUEBLO]": pueblo_centro,
        "[RESULTADO]": informe_estado,
        "[FECHA]": fecha_actual.strftime("%d/%m/%Y")

    }
    for paragraph in doc.paragraphs:
        for placeholder, valor in reemplazos.items():
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, valor)

    # Reemplazar en las tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for placeholder, valor in reemplazos.items():
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, valor)
    # Reemplazar en los encabezados (headers)
    for section in doc.sections:
        for header in section.header.paragraphs:
            for run in header.runs:
                for placeholder, valor in reemplazos.items():
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, valor)  


    # Obtener los cuadros del centro
    df_cuadros = obtener_cuadros(centro_id)

    # Buscar la tabla donde se insertarán los datos
    tabla = None
    for t in doc.tables:
        if len(t.columns) == 3:  # la tabla objetivo tiene 3 columnas
            tabla = t
            break

    if not tabla:
        raise ValueError("No se encontró una tabla con 3 columnas en la plantilla.")

    # Agregar filas con los datos
    for _, row in df_cuadros.iterrows():
        row_cells = tabla.add_row().cells
        if row['tipo'] == 'CGBT':
            row_cells[0].text = str(row['tipo'])
        elif row['numero']<=9:
            row_cells[0].text = f"{str(row['tipo'])}-0{str(row['numero'])}"
        else:
            row_cells[0].text = f"{str(row['tipo'])}-{str(row['numero'])}" 
        row_cells[1].text = row['nombre']
        row_cells[2].text = str(row['tierra_ohmnios']) if row['tierra_ohmnios'] is not None else 'N/A'

    # Guardar el documento en memoria
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)  # Mover el cursor al inicio
    file_path = "tmp/informe.docx"
    doc.save(file_path)
    with open("/tmp/informe.docx", "wb") as f:
        f.write(buffer.read())
      #output_pdf = "/tmp/informe.pdf"
      #pypandoc.download_pandoc()
      #pypandoc.convert_file("/tmp/informe.docx", to='pdf', outputfile=output_pdf)
    
    
    direccion_centro = datos_centro["direccion"] 
    cp_centro = str(int(float(datos_centro["cp"]))) if datos_centro.get("cp") else "00000"
    provincia_centro = datos_centro["provincia"]
    fname = f"{fecha_actual.strftime('%Y-%m-%d')}_{datos_centro['nombre'].split('_')[0]}_InfTierras"
    with open("/tmp/informe.docx", "rb") as pdf_file:
        st.download_button("Descargar Informe", pdf_file, file_name=(f"{fname}.docx"), mime="application/docx")

def generar_informe_word_aislamientos(centro_id):
    # Cargar la plantilla de Word
    doc = Document(PLANTILLA_AISLAMIENTOS)
    datos_centro = obtener_datos_centro(centro_id)
    nombre_centro = datos_centro.get("nombre") or "Desconocido"
    direccion_centro = datos_centro.get("direccion") or "Desconocido"
    cp_centro = str(int(float(datos_centro.get("cp")))) if datos_centro.get("cp") else "00000"
    provincia_centro = datos_centro.get("provincia") or "Desconocido"
    pueblo_centro = datos_centro.get("pueblo") or "Desconocido"
    fecha_actual = datetime.now()

    df_cuadros = obtener_cuadros(centro_id)
    medidas_aislamiento = df_cuadros["aislamiento_megaohmnios"].tolist()

    # Analizar la medición más alta y comparar con 1
    medidas_validas = [m for m in medidas_aislamiento if m is not None and m > 0]
    medida_minima = min(medidas_validas) if medidas_validas else 0
    informe_estado = "Favorable" if medida_minima >= 1 else "Desfavorable"

    reemplazos = {
        "[NOMBRE]": nombre_centro,
        "[DIRECCION]": direccion_centro,
        "[CP]": cp_centro,
        "[PROVINCIA]": provincia_centro,
        "[PUEBLO]": pueblo_centro,
        "[RESULTADO]": informe_estado,
        "[FECHA]": fecha_actual.strftime("%d/%m/%Y")

    }
    for paragraph in doc.paragraphs:
        for placeholder, valor in reemplazos.items():
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, valor)

    # Reemplazar en las tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for placeholder, valor in reemplazos.items():
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, valor)
    # Reemplazar en los encabezados (headers)
    for section in doc.sections:
        for header in section.header.paragraphs:
            for run in header.runs:
                for placeholder, valor in reemplazos.items():
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, valor)  


    # Obtener los cuadros del centro
    df_cuadros = obtener_cuadros(centro_id)

    # Buscar la tabla donde se insertarán los datos
    tabla = None
    for t in doc.tables:
        if len(t.columns) == 3:  # la tabla objetivo tiene 3 columnas
            tabla = t
            break

    if not tabla:
        raise ValueError("No se encontró una tabla con 3 columnas en la plantilla.")

    # Agregar filas con los datos
    for _, row in df_cuadros.iterrows():
        row_cells = tabla.add_row().cells
        if row['tipo'] == 'CGBT':
            row_cells[0].text = str(row['tipo'])
        elif row['numero']<=9:
            row_cells[0].text = f"{str(row['tipo'])}-0{str(row['numero'])}"
        else:
            row_cells[0].text = f"{str(row['tipo'])}-{str(row['numero'])}" 
        row_cells[1].text = row['nombre']
        row_cells[2].text = str(row['aislamiento_megaohmnios']) if row['aislamiento_megaohmnios'] is not None else 'N/A'

    # Guardar el documento en memoria
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)  # Mover el cursor al inicio
    file_path = "tmp/informe.docx"
    doc.save(file_path)
    with open("/tmp/informe.docx", "wb") as f:
        f.write(buffer.read())
      #output_pdf = "/tmp/informe.pdf"
      #pypandoc.download_pandoc()
      #pypandoc.convert_file("/tmp/informe.docx", to='pdf', outputfile=output_pdf)
    
    
    direccion_centro = datos_centro["direccion"]
    cp_centro = str(int(float(datos_centro["cp"]))) if datos_centro.get("cp") else "00000"
    provincia_centro = datos_centro["provincia"]
    fname = f"{fecha_actual.strftime('%Y-%m-%d')}_{datos_centro['nombre'].split('_')[0]}_InfAislamientos"
    with open("/tmp/informe.docx", "rb") as pdf_file:
        st.download_button("Descargar Informe", pdf_file, file_name=(f"{fname}.docx"), mime="application/docx")



PLANTILLA_BRA = "BASE_BRA.docx" 
PLANTILLA_BRA_SD = "BASE_BRA_SD.docx"

def generar_informe_word_bra(centro_id):

    meses = {
    "January": "enero", "February": "febrero", "March": "marzo", "April": "abril",
    "May": "mayo", "June": "junio", "July": "julio", "August": "agosto",
    "September": "septiembre", "October": "octubre", "November": "noviembre", "December": "diciembre"
    }

    defectos = obtener_defectos(centro_id) or []
    plantilla = PLANTILLA_BRA if defectos else PLANTILLA_BRA_SD
    doc = Document(plantilla)
    
    datos_centro = obtener_datos_centro(centro_id)
    fecha_actual = datetime.now()
    nombre_centro = datos_centro.get("nombre", "Desconocido")
    direccion_centro = datos_centro.get("direccion", "Desconocida")
    cp_centro = str(int(float(datos_centro.get("cp", 0)))) if datos_centro.get("cp") else "00000"
    provincia_centro = datos_centro.get("provincia", "Desconocida")
    pueblo_centro = datos_centro.get("pueblo", "Desconocido")
    email = datos_centro.get("email", "Desconocido")
    telf = datos_centro.get("telf", "Desconocido")
    pot = datos_centro.get("pot", "Desconocido")
    nif = datos_centro.get("nif", "Desconocido")
    cups = datos_centro.get("cups", "Desconocido")
    dia = fecha_actual.day
    mes = meses[fecha_actual.strftime("%B")] 
    año = fecha_actual.year

    reemplazos = {
        "[NOMBRE_EDIFICIO]": nombre_centro,
        "[DOMICILIO]": direccion_centro,
        "[CP]": cp_centro,
        "[PROVINCIA]": provincia_centro,
        "[MUNICIPIO]": pueblo_centro,
        "[MAIL]": email,
        "[TELEFONO]": telf,
        "[POTENCIA]": pot,
        "[NIF]": nif,
        "[CUPS]": cups,
        "[DIA]": str(dia),
        "[MES]": mes,
        "[AÑO]": str(año)
    }

    # Reemplazo en texto
    for paragraph in doc.paragraphs:
        for placeholder, valor in reemplazos.items():
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, valor)

    # Reemplazo en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for placeholder, valor in reemplazos.items():
                            if placeholder in run.text:
                                if valor is not None:
                                    run.text = run.text.replace(placeholder, valor)

    # Reemplazo en encabezado
    for section in doc.sections:
        for header in section.header.paragraphs:
            for run in header.runs:
                for placeholder, valor in reemplazos.items():
                    if placeholder in run.text:
                        if valor is not None:
                         run.text = run.text.replace(placeholder, valor)

        # Obtener la tabla de 5 columnas
    defectos = obtener_defectos(centro_id)
    print(centro_id)
    tabla = None
    for t in doc.tables:
        if len(t.columns) == 5:
            tabla = t
            break

    if not tabla:
        raise ValueError("No se encontró una tabla con 5 columnas en la plantilla.")

    cuadros_agregados = set()

    for defecto in defectos:
        cuadro = defecto["cuadro"]
        nombre_normalizado = defecto["nombre_normalizado"]
        itc = defecto["itc"]
        id= defecto["cuadro_id"]

        if cuadro not in cuadros_agregados:

            tipo, numero = obtener_datos_cuadro(id)
            row_cells = tabla.add_row().cells
            p = row_cells[0].paragraphs[0]

            if tipo == 'CGBT':
                run = p.add_run(str(tipo))
            elif numero <= 9:
                run = p.add_run(f"{tipo}-0{numero} {cuadro}")
            else:
                run = p.add_run(f"{tipo}-{numero} {cuadro}") 

            run.bold = True
            cuadros_agregados.add(cuadro)
        # Fila con defecto
        row_cells = tabla.add_row().cells
        row_cells[0].text = nombre_normalizado
        row_cells[1].text = itc
        row_cells[2].text = ""
        row_cells[3].text = "x"
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[4].text = "NO"
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
  

    # Guardar y descargar
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    file_path = "tmp/informe_bra.docx"
    doc.save(file_path)

    fname = f"{fecha_actual.strftime('%Y-%m-%d')}_{nombre_centro.split('_')[0]}_BRA"
    with open("tmp/informe_bra.docx", "rb") as docx_file:
        st.download_button("Descargar BRA", docx_file, file_name=f"{fname}.docx", mime="application/docx")        
    
def generar_informe_word_reparacion(centro_id):
    doc = Document()
    datos_centro = obtener_datos_centro(centro_id)
    nombre_centro = datos_centro.get("nombre", "Desconocido")

    # Título principal
    p = doc.add_paragraph()
    run = p.add_run(nombre_centro)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    defectos = obtener_defectos(centro_id) or []
    if not defectos:
        doc.add_paragraph("No hay defectos registrados para este centro.")
    else:
        cuadros = defaultdict(list)
        for defecto in defectos:
            cuadros[defecto["cuadro"]].append(defecto)

        df_cuadros = obtener_cuadros(centro_id)

        for cuadro, defectos_cuadro in cuadros.items():
            cuadro_p = doc.add_paragraph()

            tipo = None
            numero = None
            anotaciones = ""
            if not df_cuadros.empty:
                fila_cuadro = df_cuadros[df_cuadros["nombre"] == cuadro]
                if not fila_cuadro.empty:
                    tipo = fila_cuadro.iloc[0].get("tipo", "")
                    numero = fila_cuadro.iloc[0].get("numero", "")
                    anotaciones = fila_cuadro.iloc[0].get("anotaciones", "")

            if tipo == 'CGBT':
                cuadro_run = cuadro_p.add_run(str(cuadro))
            elif numero != "":
                numero_str = f"0{numero}" if isinstance(numero, int) and numero <= 9 else str(numero)
                cuadro_run = cuadro_p.add_run(f"{tipo}-{numero_str} {cuadro}")
            else:
                cuadro_run = cuadro_p.add_run(str(cuadro))

            cuadro_run.bold = True
            cuadro_run.font.size = Pt(16)
            cuadro_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            for defecto in defectos_cuadro:
                nombre = defecto.get("nombre_normalizado", "").strip()
                detalle = defecto.get("detalles", "").strip()

                # Mostrar siempre el nombre, y el detalle si existe
                if detalle:
                    doc.add_paragraph(f"{nombre}: {detalle}", style="List Bullet")
                else:
                    doc.add_paragraph(f"{nombre}", style="List Bullet")

            if anotaciones:
                doc.add_paragraph(f"Anotaciones:\n{anotaciones}")

    # Guardar el documento en un buffer y mostrar el botón de descarga
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    fecha_actual = datetime.now()
    fname = f"{fecha_actual.strftime('%Y-%m-%d')}_{nombre_centro.split('_')[0]}_Reparacion.docx"
    st.download_button("Descargar Informe de Reparación", buffer, file_name=fname, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
# Función para obtener el archivo Word generado
def obtener_word_tierras(centro_id):
    df_cuadros = obtener_cuadros(centro_id)
    if df_cuadros.empty:
        raise ValueError(f"No hay cuadros creados en el centro para generar el informe de tierras.")
    return generar_informe_word_tierras(centro_id)


def obtener_word_aislamientos(centro_id):
    df_cuadros = obtener_cuadros(centro_id)
    if df_cuadros.empty:
        raise ValueError(f"No hay cuadros creados en el centro para generar el informe de aislamientos.")
    return generar_informe_word_aislamientos(centro_id)


def generar_informe_bateria(centro_id, df_escalones):

    meses = {
    "January": "enero", "February": "febrero", "March": "marzo", "April": "abril",
    "May": "mayo", "June": "junio", "July": "julio", "August": "agosto",
    "September": "septiembre", "October": "octubre", "November": "noviembre", "December": "diciembre"
    }

    plantilla = "plantilla_bateria.docx" 
    doc = Document(plantilla)
    
    datos_centro = obtener_datos_centro(centro_id, "centros_bateria")
    fecha_actual = datetime.now()
    nombre_centro = datos_centro.get("nombre", "Desconocido")
    direccion_centro = datos_centro.get("direccion", "Desconocida")
    cp_centro = str(int(float(datos_centro.get("cp", 0)))) if datos_centro.get("cp") else "00000"
    provincia_centro = datos_centro.get("provincia", "Desconocida")
    pueblo_centro = datos_centro.get("pueblo", "Desconocido")
    marcar = datos_centro.get("marca_regulador", "Desconocido")
    tipor = datos_centro.get("tipo_regulador", "Desconocido")
    marcac = datos_centro.get("marca_condensadores", "Desconocido")
    modeloc = datos_centro.get("modelo_condensadores", "Desconocido")
    tension = datos_centro.get("tension_servicio", "Desconocido")
    potenciac = datos_centro.get("potencia_condensadores", "Desconocido")
    escalones = str(datos_centro.get("num_escalones", "Desconocido"))
    potenciat = datos_centro.get("potencia_total", "Desconocido")
    seccionl= datos_centro.get("seccion_linea", "Desconocido")
    estadov = datos_centro.get("estado_visual", "Desconocido")
    referencia = datos_centro.get("referencia_equipo", "Desconocido")
    comentario = datos_centro.get("comentario", "No hay comentarios")
    dia = fecha_actual.day
    mes = meses[fecha_actual.strftime("%B")] 
    año = fecha_actual.year

    reemplazos = {
        "[NOMBRE_EDIFICIO]": nombre_centro,
        "[DOMICILIO]": direccion_centro,
        "[CP]": cp_centro,
        "[PROVINCIA]": provincia_centro,
        "[MUNICIPIO]": pueblo_centro,
        "[MARCAR]": marcar,
        "[TIPOR]": tipor,
        "[MARCAC]": marcac,
        "[MODELOC]": modeloc,
        "[TENSION]": tension,
        "[POTENCIA]": potenciac,    
        "[ESCALONES]": escalones,
        "[QTOT]": potenciat,
        "[SECC]": seccionl, 
        "[VISUAL]": estadov,
        "[REF]": referencia,
        "[DIA]": str(dia),
        "[MES]": mes,
        "[AÑO]": str(año),
        "[COMENTARIO]": comentario

    }

    # Reemplazo en texto
    for paragraph in doc.paragraphs:
        for placeholder, valor in reemplazos.items():
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, valor)

    # Reemplazo en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for placeholder, valor in reemplazos.items():
                            if placeholder in run.text:
                                if valor is not None:
                                    run.text = run.text.replace(placeholder, valor)

    # Reemplazo en encabezado
    for section in doc.sections:
        for header in section.header.paragraphs:
            for run in header.runs:
                for placeholder, valor in reemplazos.items():
                    if placeholder in run.text:
                        if valor is not None:
                         run.text = run.text.replace(placeholder, valor)

    tabla = None
    for tbl in doc.tables:
        if "POTENCIA NOMINAL" in tbl.cell(0, 1).text:
            tabla = tbl
            break

    if tabla is None:
        raise ValueError("❌ No se encontró la tabla de mediciones en el documento.")

    # Recorremos los escalones y rellenamos las filas (comienza en fila 1, fila 0 es cabecera)
    for i in range(min(len(df_escalones), len(tabla.rows) - 1)):
        fila = tabla.rows[i + 2]
        fila.cells[1].text = str(df_escalones.iloc[i]["POTENCIA NOMINAL (kVAr)"])
        fila.cells[2].text = str(df_escalones.iloc[i]["INTENSIDAD NOMINAL (A)"])
        fila.cells[3].text = str(df_escalones.iloc[i]["CONSUMO R (A)"])
        fila.cells[4].text = str(df_escalones.iloc[i]["CONSUMO S (A)"])
        fila.cells[5].text = str(df_escalones.iloc[i]["CONSUMO T (A)"])
        fila.cells[6].text = str(df_escalones.iloc[i]["RENDIMIENTO R (%)"])
        fila.cells[7].text = str(df_escalones.iloc[i]["RENDIMIENTO S (%)"])
        fila.cells[8].text = str(df_escalones.iloc[i]["RENDIMIENTO T (%)"])


    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    file_path = "tmp/informe_batcond.docx"
    doc.save(file_path)

    filename = f"{fecha_actual.strftime('%Y-%m-%d')}_{nombre_centro}_BatCond"
    filename = re.sub(r"[\r\n]", "", filename)  # elimina saltos de línea
    filename = filename.replace("/", "-")  # opcional: evita errores por nombres con slashes
    with open("tmp/informe_batcond.docx", "rb") as docx_file:
        st.download_button("Descargar Informe", docx_file, file_name=f"{filename}.docx", mime="application/docx")

def generar_presupuesto(centro_id):

    defectos_presupuestables = {
    "PEGATINA": {
        "Articulo": "BT001",
        "Denominación": "Suministro y colocación de placa de señalización de riesgo eléctrico en el cuadro",
        "Cantidad": 1
    },
    "PUNTERAS": {
        "Articulo": "BT012",
        "Denominación": "Trabajos de desconexión del cableado que se encuentra sin punteras y/o mal conexionado y realizar el correcto conexionado con punteras.",
        "Cantidad": 1
    },
    "CIR SIN IDENTIFICAR": {
        "Articulo": "BT004",
        "Denominación": "Trabajos de identificación y marcado de los circuitos del cuadro que se encuentren sin identificar.",
        "Cantidad": 1
    },
    "IDENTIF. COLORES": {
        "Articulo": "BT003",
        "Denominación": "Trabajos de identificación de los conductores mediante colores que se encuentran sin identificar.",
        "Cantidad": 1
    },
    "POLARIDAD INVERTIDA": {
        "Articulo": "BT002",
        "Denominación": "Trabajos de desconexión del circuito conectado aguas arriba y conexión.",
        "Cantidad": 1
    },
    "OBTURADORES": {
        "Articulo": "BT014",
        "Denominación": "Suministro y montaje de obturadores.",
        "Cantidad": 1
    },
    "PUERTAS/CHASIS": {
        "Articulo": "BT009",
        "Denominación": "Trabajos de puesta a tierra directa del cuadro electrico.",
        "Cantidad": 1
    }
    }

    df = obtener_cuadros(centro_id)
    columnas = [
        "Partida", "SubPartida", "Articulo", "Denominación","Cantidad"
    ]

    filas = []
    filas.append({
        "Partida": "Instalación Eléctrica",
        "Denominación": "Instalación Eléctrica"
    })

    for idx, (_, row) in enumerate(df.iterrows(), start= 1):
        defectos_celda = row.get("defectos", "")
        if defectos_celda:  # Solo si la columna defectos no está vacía
            filas.append({
                "SubPartida": f"1. {idx}.",
                "Denominación": row["nombre"],
                "Cantidad"  : 1
            })
        defectos_celda = row.get("defectos", "")

# Convertimos a lista de defectos estandarizada
        if isinstance(defectos_celda, list) or isinstance(defectos_celda, tuple):
            lista_defectos = [d.strip().upper() for d in defectos_celda]
        elif isinstance(defectos_celda, str) and defectos_celda.strip() != "":
            lista_defectos = [d.strip().upper() for d in defectos_celda.split(",")]
        else:
            lista_defectos = []
        print(f"Lista de defectos: {lista_defectos}")
        # Recorremos los defectos encontrados
        for defecto in lista_defectos:
            if defecto in defectos_presupuestables:
                info = defectos_presupuestables[defecto]
                filas.append({
                    "Articulo": info["Articulo"],
                    "Denominación": info["Denominación"],
                    "Cantidad": info["Cantidad"]
                })

    df_out = pd.DataFrame(filas, columns=columnas)
    df_out.to_excel("presupuesto.xlsx", index=False)   

    with open("presupuesto.xlsx", "rb") as f:
        bytes_data = f.read()
    
    st.download_button(
        label="Descargar Excel",
        data=bytes_data,
        file_name="presupuesto.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

def generar_informe_word_castellon(centro_id):

    meses = {
    "January": "enero", "February": "febrero", "March": "marzo", "April": "abril",
    "May": "mayo", "June": "junio", "July": "julio", "August": "agosto",
    "September": "septiembre", "October": "octubre", "November": "noviembre", "December": "diciembre"
    }
    PLANTILLA_CASTELLON = "BASE_CASTELLON.docx"
    defectos = obtener_defectos(centro_id) or []
    plantilla = PLANTILLA_CASTELLON
    doc = Document(plantilla)
    
    datos_centro = obtener_datos_centro(centro_id)
    fecha_actual = datetime.now()
    nombre_centro = datos_centro.get("nombre", "Desconocido")
    direccion_centro = datos_centro.get("direccion", "Desconocida")
    cp_centro = str(int(float(datos_centro.get("cp", 0)))) if datos_centro.get("cp") else "00000"
    provincia_centro = datos_centro.get("provincia", "Desconocida")
    pueblo_centro = datos_centro.get("pueblo", "Desconocido")
    email = datos_centro.get("email", "Desconocido")
    telf = datos_centro.get("telf", "Desconocido")
    pot = datos_centro.get("pot", "Desconocido")
    nif = datos_centro.get("nif", "Desconocido")
    cups = datos_centro.get("cups", "Desconocido")
    dia = fecha_actual.day
    mes = meses[fecha_actual.strftime("%B")] 
    año = fecha_actual.year

    reemplazos = {
        "[NOMBRE_EDIFICIO]": nombre_centro,
        "[DOMICILIO]": direccion_centro,
        "[CP]": cp_centro,
        "[PROVINCIA]": provincia_centro,
        "[MUNICIPIO]": pueblo_centro,
        "[MAIL]": email,
        "[TELEFONO]": telf,
        "[POTENCIA]": pot,
        "[NIF]": nif,
        "[CUPS]": cups,
        "[DIA]": str(dia),
        "[MES]": mes,
        "[AÑO]": str(año)
    }

    # Reemplazo en texto
    for paragraph in doc.paragraphs:
        for placeholder, valor in reemplazos.items():
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, valor)

    # Reemplazo en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for placeholder, valor in reemplazos.items():
                            if placeholder in run.text:
                                if valor is not None:
                                    run.text = run.text.replace(placeholder, valor)

    # Reemplazo en encabezado
    for section in doc.sections:
        for header in section.header.paragraphs:
            for run in header.runs:
                for placeholder, valor in reemplazos.items():
                    if placeholder in run.text:
                        if valor is not None:
                         run.text = run.text.replace(placeholder, valor)

        # Obtener la tabla de 5 columnas
    df_cuadros = obtener_cuadros(centro_id)

    # Buscar la tabla donde se insertarán los datos
    tabla = None
    for t in doc.tables:
        if len(t.columns) == 5:  # la tabla objetivo tiene 5 columnas
            tabla = t
            break

    if not tabla:
        raise ValueError("No se encontró una tabla con 5 columnas en la plantilla.")

    # Agregar filas con los datos
    for _, row in df_cuadros.iterrows():
        row_cells = tabla.add_row().cells
        if row['tipo'] == 'CGBT':
            row_cells[0].text = str(row['tipo'])
        elif row['numero']<=9:
            row_cells[0].text = f"{str(row['tipo'])}-0{str(row['numero'])}"
        else:
            row_cells[0].text = f"{str(row['tipo'])}-{str(row['numero'])}" 
        row_cells[1].text = row['nombre']
        row_cells[2].text = str(row['tierra_ohmnios']) if row['tierra_ohmnios'] is not None else 'N/A'
        row_cells[4].text = str(row['aislamiento_megaohmnios']) if row['aislamiento_megaohmnios'] is not None else 'N/A'

    def set_cell_bg_color(cell, color_hex):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)

    # Función para formatear texto (negrita, tamaño)
    def format_text(cell, bold=False, font_size=None):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = bold
                if font_size:
                    run.font.size = Pt(font_size)

    # Función para calcular calificación
    itc_leves = {"R.E.B.T. Art. 6", "RD 485/1997"} 

    def calcular_calificacion(defecto):
        if defecto["itc"] in itc_leves:
            return "leve"
        return "grave"
    
    # Agrupar defectos por cuadro
    defectos_por_cuadro = {}
    for defecto in defectos:
        cuadro = defecto["cuadro"]
        if cuadro not in defectos_por_cuadro:
            defectos_por_cuadro[cuadro] = []
        defectos_por_cuadro[cuadro].append(defecto)

    # Recorrer cada cuadro y generar su tabla de título y su tabla de defectos
    total_cuadros = len(defectos_por_cuadro)
    for idx, (cuadro, lista_defectos) in enumerate(defectos_por_cuadro.items(), start=1):

        tipo, numero = obtener_datos_cuadro(lista_defectos[0]["cuadro_id"])

        # Determinar código del cuadro
        if tipo == 'CGBT':
            codigo_cuadro = tipo
        elif numero <= 9:
            codigo_cuadro = f"{tipo}-0{numero}"
        else:
            codigo_cuadro = f"{tipo}-{numero}"

        denominacion = cuadro

        # Tabla del título (2x2)
        tabla_titulo = doc.add_table(rows=2, cols=2)
        tabla_titulo.style = 'Table Grid'
        tabla_titulo.cell(0, 0).text = "Código del cuadro"
        tabla_titulo.cell(0, 1).text = codigo_cuadro
        tabla_titulo.cell(1, 0).text = "Denominación"
        tabla_titulo.cell(1, 1).text = denominacion
        tabla_titulo.columns[0].width = Pt(1)
        tabla_titulo.columns[1].width = Pt(2)
        # Negrita en las etiquetas de la tabla título
        format_text(tabla_titulo.cell(0, 0), bold=True)
        format_text(tabla_titulo.cell(1, 0), bold=True)

        set_cell_bg_color(tabla_titulo.cell(0, 0), "a7d08c")
        set_cell_bg_color(tabla_titulo.cell(1, 0), "a7d08c")
        format_text(tabla_titulo.cell(0, 0), bold=True)
        format_text(tabla_titulo.cell(1, 0), bold=True)

        doc.add_paragraph("")  # Espacio después del título

        # Crear tabla de defectos
        tabla = doc.add_table(rows=1, cols=3)
        tabla.style = 'Table Grid'
        hdr_cells = tabla.rows[0].cells
        hdr_cells[0].text = "Descripción"
        hdr_cells[1].text = "Normativa"
        hdr_cells[2].text = "Calificación"

        # Formato de encabezado
        for cell in hdr_cells:
            set_cell_bg_color(cell, "a7d08c")
            format_text(cell, bold=True, font_size=11)  # Encabezado 1 punto más grande

        # Ajustar anchos → descripción doble de ancho que los demás
        tabla.columns[0].width = Pt(900)  # más ancho
        tabla.columns[1].width = Pt(35)
        tabla.columns[2].width = Pt(35)

        # Añadir filas con defectos
        for defecto in lista_defectos:
            row_cells = tabla.add_row().cells
            row_cells[0].text = defecto["nombre_normalizado"]
            row_cells[1].text = defecto["itc"]
            row_cells[2].text = calcular_calificacion(defecto).capitalize()

        # Salto de página entre cuadros
        if idx < total_cuadros:
            doc.add_page_break()
            # Guardar y descargar
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            file_path = "tmp/informe_castellon.docx"
            doc.save(file_path)

    fname = f"{fecha_actual.strftime('%Y-%m-%d')}_{nombre_centro.split('_')[0]}_RevBT"
    with open("tmp/informe_castellon.docx", "rb") as docx_file:
        st.download_button(
        "Descargar Informe Extendido",
        docx_file,
        file_name=f"{fname}.docx",
        mime="application/docx",
        key=f"descargar_{fname}_{uuid.uuid4()}"
    )    